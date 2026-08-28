"""Unit tests for document parsers."""

from __future__ import annotations

import io

import pytest

from app.services.ingestion import parsers
from app.services.ingestion.parsers import UnsupportedFileTypeError, parse


def test_parse_txt() -> None:
    doc = parse(filename="notes.txt", content_type="text/plain", data=b"hello   world")
    assert doc.pages[0].page is None
    assert doc.full_text == "hello world"


def test_parse_markdown_normalizes_blank_lines() -> None:
    data = b"# Title\n\n\n\nBody   text"
    doc = parse(filename="readme.md", content_type="text/markdown", data=data)
    assert "\n\n\n" not in doc.full_text
    assert "Body text" in doc.full_text


def test_dispatch_falls_back_to_mime_type() -> None:
    # No useful extension, but the MIME type identifies it as text.
    doc = parse(filename="upload", content_type="text/plain", data=b"content")
    assert doc.full_text == "content"


def test_unsupported_type_raises() -> None:
    with pytest.raises(UnsupportedFileTypeError):
        parse(filename="image.png", content_type="image/png", data=b"\x89PNG")


def test_parse_docx_roundtrip() -> None:
    from docx import Document as DocxDocument

    docx_doc = DocxDocument()
    docx_doc.add_paragraph("First paragraph.")
    docx_doc.add_paragraph("Second paragraph.")
    buf = io.BytesIO()
    docx_doc.save(buf)

    parsed = parse(
        filename="doc.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        data=buf.getvalue(),
    )
    assert "First paragraph." in parsed.full_text
    assert "Second paragraph." in parsed.full_text


def test_parse_pdf_maps_pages(monkeypatch: pytest.MonkeyPatch) -> None:
    class FakePage:
        def __init__(self, text: str) -> None:
            self._text = text

        def extract_text(self) -> str:
            return self._text

    class FakeReader:
        def __init__(self, _stream: object) -> None:
            self.pages = [FakePage("Page one text"), FakePage("Page two text")]

    monkeypatch.setattr("pypdf.PdfReader", FakeReader)

    parsed = parse(filename="report.pdf", content_type="application/pdf", data=b"%PDF-fake")
    assert [p.page for p in parsed.pages] == [1, 2]
    assert parsed.pages[0].text == "Page one text"


def test_pdf_with_no_text_raises(monkeypatch: pytest.MonkeyPatch) -> None:
    class FakeReader:
        def __init__(self, _stream: object) -> None:
            self.pages = [type("P", (), {"extract_text": lambda self: ""})()]

    monkeypatch.setattr("pypdf.PdfReader", FakeReader)

    with pytest.raises(UnsupportedFileTypeError):
        parse(filename="scan.pdf", content_type="application/pdf", data=b"%PDF")


def test_supported_extensions_exposed() -> None:
    assert ".pdf" in parsers.SUPPORTED_EXTENSIONS
    assert ".docx" in parsers.SUPPORTED_EXTENSIONS
