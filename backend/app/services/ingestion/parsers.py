"""Document parsers: raw bytes -> normalized text with provenance.

Dispatch is by file extension first (most reliable), then MIME type. Each
parser returns a ``ParsedDocument`` composed of ``ParsedPage`` items so that
page-level provenance (for PDFs) survives into chunk metadata and citations.
"""

from __future__ import annotations

import io
import re
from pathlib import Path

from pydantic import BaseModel

from app.core.errors import AppError
from app.core.logging import get_logger

logger = get_logger(__name__)

_WHITESPACE_RE = re.compile(r"[ \t]+")
_BLANKLINES_RE = re.compile(r"\n{3,}")


class UnsupportedFileTypeError(AppError):
    status_code = 422
    code = "unsupported_file_type"


class ParsedPage(BaseModel):
    page: int | None
    text: str


class ParsedDocument(BaseModel):
    pages: list[ParsedPage]

    @property
    def full_text(self) -> str:
        return "\n\n".join(p.text for p in self.pages if p.text).strip()


def _normalize(text: str) -> str:
    """Collapse horizontal whitespace and excess blank lines; strip edges."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _WHITESPACE_RE.sub(" ", text)
    text = _BLANKLINES_RE.sub("\n\n", text)
    return text.strip()


def _parse_text(data: bytes) -> ParsedDocument:
    try:
        raw = data.decode("utf-8")
    except UnicodeDecodeError:
        raw = data.decode("latin-1")
    return ParsedDocument(pages=[ParsedPage(page=None, text=_normalize(raw))])


def _parse_pdf(data: bytes) -> ParsedDocument:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages: list[ParsedPage] = []
    for i, page in enumerate(reader.pages, start=1):
        text = _normalize(page.extract_text() or "")
        if text:
            pages.append(ParsedPage(page=i, text=text))
    if not pages:
        # A scanned/image PDF yields no extractable text.
        raise UnsupportedFileTypeError("No extractable text found in PDF (is it scanned?)")
    return ParsedDocument(pages=pages)


def _parse_docx(data: bytes) -> ParsedDocument:
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs)
    return ParsedDocument(pages=[ParsedPage(page=None, text=_normalize(text))])


# extension -> parser
_PARSERS = {
    ".pdf": _parse_pdf,
    ".docx": _parse_docx,
    ".txt": _parse_text,
    ".md": _parse_text,
    ".markdown": _parse_text,
}

# mime type -> extension (fallback when the filename has no useful suffix)
_MIME_FALLBACK = {
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "text/plain": ".txt",
    "text/markdown": ".md",
}

SUPPORTED_EXTENSIONS = frozenset(_PARSERS)


def parse(*, filename: str, content_type: str, data: bytes) -> ParsedDocument:
    """Parse an uploaded file into a normalized ``ParsedDocument``.

    Raises ``UnsupportedFileTypeError`` for types we cannot handle.
    """
    ext = Path(filename).suffix.lower()
    if ext not in _PARSERS:
        ext = _MIME_FALLBACK.get(content_type, ext)

    parser = _PARSERS.get(ext)
    if parser is None:
        raise UnsupportedFileTypeError(
            f"Unsupported file type: {filename!r} ({content_type}). "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    parsed = parser(data)
    logger.info("parsed_document", filename=filename, pages=len(parsed.pages))
    return parsed
